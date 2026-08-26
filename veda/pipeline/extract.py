"""Read uploaded project documents (spec 33) into text/rows, then into evidence.

Supports the paperwork a construction project actually produces: CSV and Excel
trackers, PDFs, Word documents, JSON, plain text, chat exports and site reports.
Extraction is deterministic and rule-based; interpretation is the agent's job.
"""
from __future__ import annotations

import csv
import io
import json
import os
import re
from datetime import datetime
from typing import Any

from .. import db

TABULAR = {".csv", ".tsv"}
EXCEL = {".xlsx", ".xlsm", ".xls"}
TEXTY = {".txt", ".md", ".log", ".json"}


# ------------------------------------------------------------------ reading
def read_rows(path: str, ext: str, limit: int = 100_000) -> tuple:
    """Return (headers, rows, sheet_names) for tabular sources."""
    ext = ext.lower()
    if ext in TABULAR:
        with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
            sample = f.read(8192)
            f.seek(0)
            delim = "\t" if ext == ".tsv" else _sniff(sample)
            rdr = csv.reader(f, delimiter=delim)
            all_rows = []
            for i, r in enumerate(rdr):
                if i > limit:
                    break
                all_rows.append(r)
        if not all_rows:
            return [], [], []
        return all_rows[0], all_rows[1:], ["(csv)"]

    if ext in EXCEL:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        headers: list = []
        rows: list = []
        sheets = []
        for ws in wb.worksheets:
            sheets.append(ws.title)
            first = True
            for r in ws.iter_rows(values_only=True):
                vals = ["" if v is None else str(v) for v in r]
                if not any(v.strip() for v in vals):
                    continue
                if first:
                    if not headers:
                        headers = vals
                    first = False
                    continue
                rows.append(vals + ["__sheet__:" + ws.title])
                if len(rows) > limit:
                    break
        wb.close()
        return headers, rows, sheets
    return [], [], []


def _sniff(sample: str) -> str:
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except Exception:
        return ","


def read_text(path: str, ext: str, limit: int = 2_000_000) -> str:
    ext = ext.lower()
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            return ""
        try:
            rd = PdfReader(path)
            out = []
            for i, pg in enumerate(rd.pages):
                try:
                    t = pg.extract_text() or ""
                except Exception:
                    t = ""
                out.append("[page " + str(i + 1) + "]\n" + t)
                if sum(len(x) for x in out) > limit:
                    break
            return "\n\n".join(out)
        except Exception:
            return ""
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            return ""
        try:
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for tb in d.tables:
                for row in tb.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)[:limit]
        except Exception:
            return ""
    if ext in TEXTY or ext in TABULAR:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read(limit)
        except Exception:
            return ""
    return ""


def full_text(path: str, ext: str) -> str:
    """Text used for the security scan and for free-text evidence."""
    ext = ext.lower()
    if ext in EXCEL:
        headers, rows, _ = read_rows(path, ext, limit=5000)
        buf = [" | ".join(headers)]
        for r in rows[:5000]:
            buf.append(" | ".join(str(x) for x in r))
        return "\n".join(buf)
    return read_text(path, ext)


def read_for_agent(f: dict, offset: int = 0, limit: int = 200) -> dict:
    """Paged read used by the veda_read_file MCP tool."""
    path, ext = f["stored_path"], (f.get("ext") or "").lower()
    if ext in TABULAR or ext in EXCEL:
        headers, rows, sheets = read_rows(path, ext)
        window = rows[offset:offset + max(1, min(limit, 500))]
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(headers)
        for r in window:
            w.writerow(r)
        return {"file": f.get("filename"), "kind": "tabular", "sheets": sheets,
                "total_rows": len(rows), "offset": offset,
                "returned": len(window), "content": buf.getvalue()}
    text = read_text(path, ext)
    chars = max(1000, min(limit, 200) * 1000)
    window = text[offset:offset + chars]
    return {"file": f.get("filename"), "kind": "text", "total_chars": len(text),
            "offset": offset, "returned": len(window), "content": window}


# --------------------------------------------------------- evidence mapping
# Column synonyms seen across DPRs, registers and trackers.
FIELD_MAP = {
    "date": ["date", "report_date", "dpr_date", "day", "reporting_date", "doc_date"],
    "author": ["reported_by", "author", "prepared_by", "raised_by", "inspector",
               "reporter", "supervisor"],
    "contractor": ["contractor", "subcontractor", "agency", "vendor", "supplier",
                   "firm"],
    "crew": ["crew", "crew_id", "gang", "team", "shift_crew", "welder_id"],
    "discipline": ["discipline", "trade", "activity_type", "work_type", "category",
                   "ndt_method"],
    "location": ["location", "spread", "area", "zone", "section", "site", "unit"],
    "chainage": ["chainage", "chainage_from", "ch", "km", "station", "kp",
                 "chainage_start"],
    "quantity": ["quantity", "qty", "progress_qty", "achieved", "output", "value",
                 "no_of_joints", "count"],
    "unit": ["unit", "uom", "units", "measure"],
    "description": ["description", "remarks", "activity", "work_done", "narrative",
                    "subject", "comments", "observation", "detail", "scope"],
    "progress": ["progress_pct", "progress", "percent_complete", "pct", "completion",
                 "percentage"],
    "ref": ["dpr_no", "ncr_no", "id", "ref", "reference", "weld_no", "no", "sr_no",
            "mrn_no", "si_no", "doc_no"],
    "status": ["status", "state", "ndt_result", "result", "disposition"],
}

DISCIPLINE_HINTS = {
    "weld": "Welding", "ndt": "NDT", "radiograph": "NDT", "rt": "NDT",
    "coat": "Coating", "trench": "Trenching", "string": "Stringing",
    "bend": "Bending", "lower": "Lowering-in", "backfill": "Backfilling",
    "hydro": "Hydrotest", "test": "Testing", "civil": "Civil",
    "electric": "E&I", "instrument": "E&I", "hdd": "HDD", "crossing": "Crossing",
    "tie-in": "Tie-in", "tie in": "Tie-in", "commission": "Commissioning",
}

_DATE_FORMATS = ["%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%d.%m.%Y",
                 "%Y/%m/%d", "%d %b %Y", "%d %B %Y", "%b %d, %Y", "%Y-%m-%d %H:%M:%S"]


def norm_date(v: Any) -> str | None:
    if v in (None, ""):
        return None
    s = str(v).strip()
    if not s:
        return None
    s = s.split(" ")[0] if re.match(r"^\d{4}-\d{2}-\d{2}[ T]", s) else s
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            continue
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        return m.group(0)
    m = re.search(r"(\d{2})/(\d{2})/(\d{4})", s)
    if m:
        try:
            return datetime.strptime(m.group(0), "%d/%m/%Y").date().isoformat()
        except ValueError:
            return None
    return None


def _norm_key(h: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(h or "").strip().lower()).strip("_")


def build_header_index(headers: list) -> dict:
    """Map canonical field -> column position, using synonyms."""
    idx = {}
    normed = [_norm_key(h) for h in headers]
    for field, names in FIELD_MAP.items():
        for i, h in enumerate(normed):
            if h in names:
                idx.setdefault(field, i)
        if field not in idx:
            for i, h in enumerate(normed):
                if any(h.startswith(n) or n in h for n in names):
                    idx.setdefault(field, i)
                    break
    return idx


def _get(row: list, idx: dict, field: str) -> str:
    i = idx.get(field)
    if i is None or i >= len(row):
        return ""
    v = row[i]
    return "" if v is None else str(v).strip()


def _num(v: str):
    if not v:
        return None
    m = re.search(r"-?\d+(?:\.\d+)?", str(v).replace(",", ""))
    return float(m.group(0)) if m else None


def guess_discipline(*texts: str) -> str | None:
    blob = " ".join(t for t in texts if t).lower()
    for k, v in DISCIPLINE_HINTS.items():
        if k in blob:
            return v
    return None


def extract_evidence(project_id: str, f: dict, job_id: str | None = None) -> list:
    """Deterministic evidence extraction from one file (spec 34).

    Rows become evidence items with provenance SOURCE_FILE. No interpretation
    happens here - linking and judgement are separate, reviewable steps.
    """
    ext = (f.get("ext") or "").lower()
    path = f["stored_path"]
    out: list = []

    if ext in TABULAR or ext in EXCEL:
        headers, rows, _ = read_rows(path, ext)
        if not headers:
            return out
        idx = build_header_index(headers)
        heads = [_norm_key(h) for h in headers]
        for n, row in enumerate(rows, start=2):
            sheet = None
            if row and str(row[-1]).startswith("__sheet__:"):
                sheet = str(row[-1]).split(":", 1)[1]
                row = row[:-1]
            if not any(str(c).strip() for c in row):
                continue
            desc = _get(row, idx, "description")
            ref = _get(row, idx, "ref")
            status = _get(row, idx, "status")
            if not desc:
                # Build a readable line from whatever the row does carry.
                parts = []
                for i, h in enumerate(heads):
                    if i < len(row) and str(row[i]).strip() and h not in (
                            "date", "crew", "contractor", "unit"):
                        parts.append(str(headers[i]) + "=" + str(row[i]).strip())
                desc = "; ".join(parts[:8])
            disc = _get(row, idx, "discipline") or guess_discipline(
                desc, f.get("filename", ""), ref)
            loc = _get(row, idx, "location")
            locator = ("sheet " + sheet + ", row " + str(n)) if sheet \
                else ("row " + str(n))
            item = {
                "project_id": project_id, "file_id": f["id"], "job_id": job_id,
                "source_file": f.get("filename"), "locator": locator,
                "date": norm_date(_get(row, idx, "date")),
                "author": _get(row, idx, "author") or None,
                "contractor": _get(row, idx, "contractor") or None,
                "crew": _get(row, idx, "crew") or None,
                "discipline": disc or None,
                "location": loc or None,
                "chainage": _get(row, idx, "chainage") or None,
                "quantity": _num(_get(row, idx, "quantity")),
                "unit": _get(row, idx, "unit") or None,
                "description": (ref + ": " if ref else "") + desc[:900],
                "observed_progress": _num(_get(row, idx, "progress")),
                "confidence": 0.55,
                "state": "new",
                "security_state": f.get("security_state", "clean"),
                "raw_json": db.jdumps(dict(zip(headers, row))),
                "provenance": "SOURCE_FILE",
            }
            if status:
                item["raw_json"] = db.jdumps({**db.jloads(item["raw_json"], {}),
                                              "_status": status})
            out.append(item)
        return out

    # Free text: paragraphs and chat lines become individually citable items.
    text = read_text(path, ext)
    if not text.strip():
        return out
    if ext == ".json":
        return _extract_json(project_id, f, job_id, text)

    chat = re.findall(r"^\[([^\]]+)\]\s*([^:]{1,60}):\s*(.+)$", text, re.M)
    if len(chat) >= 3:
        for i, (stamp, who, msg) in enumerate(chat, start=1):
            out.append({
                "project_id": project_id, "file_id": f["id"], "job_id": job_id,
                "source_file": f.get("filename"), "locator": "message " + str(i),
                "date": norm_date(stamp.split(",")[0]),
                "author": who.strip() or None,
                "discipline": guess_discipline(msg),
                "description": msg.strip()[:900],
                "confidence": 0.45, "state": "new",
                "security_state": f.get("security_state", "clean"),
                "raw_json": db.jdumps({"timestamp": stamp, "from": who, "text": msg}),
                "provenance": "SOURCE_FILE",
            })
        return out

    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if len(b.strip()) > 60]
    for i, b in enumerate(blocks, start=1):
        head = b.split("\n", 1)[0][:120]
        out.append({
            "project_id": project_id, "file_id": f["id"], "job_id": job_id,
            "source_file": f.get("filename"),
            "locator": _locator_for(b, i),
            "date": norm_date(_first_date(b)),
            "discipline": guess_discipline(b),
            "description": b[:900],
            "confidence": 0.4, "state": "new",
            "security_state": f.get("security_state", "clean"),
            "raw_json": db.jdumps({"heading": head}),
            "provenance": "SOURCE_FILE",
        })
    return out


def _locator_for(block: str, i: int) -> str:
    m = re.search(r"\[page (\d+)\]", block)
    if m:
        return "page " + m.group(1)
    return "section " + str(i)


def _first_date(text: str) -> str | None:
    m = re.search(r"\d{4}-\d{2}-\d{2}|\d{1,2}[/.]\d{1,2}[/.]\d{4}|"
                  r"\d{1,2}\s+\w+\s+\d{4}", text)
    return m.group(0) if m else None


def _extract_json(project_id: str, f: dict, job_id, text: str) -> list:
    try:
        data = json.loads(text)
    except Exception:
        return []
    rows = data if isinstance(data, list) else \
        next((v for v in data.values() if isinstance(v, list)), [])
    out = []
    for i, r in enumerate(rows[:5000], start=1):
        if not isinstance(r, dict):
            continue
        low = {_norm_key(k): v for k, v in r.items()}
        def pick(field):
            for n in FIELD_MAP[field]:
                if n in low and low[n] not in (None, ""):
                    return str(low[n])
            return ""
        out.append({
            "project_id": project_id, "file_id": f["id"], "job_id": job_id,
            "source_file": f.get("filename"), "locator": "item " + str(i),
            "date": norm_date(pick("date")),
            "author": pick("author") or None,
            "contractor": pick("contractor") or None,
            "crew": pick("crew") or None,
            "discipline": pick("discipline") or guess_discipline(db.jdumps(r)),
            "location": pick("location") or None,
            "chainage": pick("chainage") or None,
            "quantity": _num(pick("quantity")),
            "unit": pick("unit") or None,
            "description": (pick("description") or db.jdumps(r))[:900],
            "observed_progress": _num(pick("progress")),
            "confidence": 0.5, "state": "new",
            "security_state": f.get("security_state", "clean"),
            "raw_json": db.jdumps(r), "provenance": "SOURCE_FILE",
        })
    return out
