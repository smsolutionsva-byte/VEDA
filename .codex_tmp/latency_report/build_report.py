from __future__ import annotations

from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Red\Downloads\VEDA-main")
WORK = ROOT / ".codex_tmp" / "latency_report"
OUT_DIR = ROOT / "reports"
OUT = OUT_DIR / "VEDA_Latency_Optimization_Report_Shivansh_Mukhia.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "2A9D8F"
GREEN = "15803D"
GOLD = "B98900"
RED = "9B1C1C"
INK = "1F2937"
MUTED = "64748B"
PALE_BLUE = "E8EEF5"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF6EF"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run(run, *, size=None, bold=None, color=INK, italic=None,
            font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border_bottom(paragraph, color=BLUE, size=12, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    keep.set(qn("w:val"), "1" if value else "0")


def set_picture_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_document_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_running_furniture(section):
    fp = section.footer.paragraphs[0]
    add_page_number(fp)


def add_text(doc, text, *, bold_lead=None, color=INK, size=11,
             after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), size=size, bold=True, color=color)
        set_run(p.add_run(text[len(bold_lead):]), size=size, color=color)
    else:
        set_run(p.add_run(text), size=size, color=color)
    return p


def add_callout(doc, label, text, *, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=180)
    set_repeat_table_header(table.rows[0])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(label + "  "), size=10.5, bold=True, color=accent)
    set_run(p.add_run(text), size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_figure(doc, path, caption, alt, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_picture_alt(shape, caption, alt)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_with_next(cp, False)


def add_table(doc, headers, rows, widths, *, header_fill=LIGHT,
              aligns=None, font_size=9.2, cell_tb=80):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = (aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
        set_run(p.add_run(header), size=9, bold=True, color=NAVY)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        if r_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = (aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths, indent=120)
    if cell_tb != 80:
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=cell_tb, bottom=cell_tb,
                                 start=120, end=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _font(size, bold=False):
    path = Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf")
    return ImageFont.truetype(str(path), size)


def _chart_canvas(width, height, title):
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 42), title, font=_font(32, True), fill="#0B2545")
    return image, draw


def create_charts():
    # Chart 1: measured resolver comparison.
    image, draw = _chart_canvas(1600, 560, "Measured resolver latency")
    x0, x1, top, max_v = 390, 1470, 150, 175
    for tick in (0, 50, 100, 150):
        x = x0 + int((x1 - x0) * tick / max_v)
        draw.line((x, top - 20, x, 430), fill="#E2E8F0", width=2)
        draw.text((x - 16, 445), str(tick), font=_font(22), fill="#64748B")
    for y, label, value, color in (
        (185, "Before optimization", 157.5, "#CBD5E1"),
        (315, "After optimization", 22.8, "#2E74B5"),
    ):
        draw.text((65, y + 16), label, font=_font(25, True), fill="#334155")
        width = int((x1 - x0) * value / max_v)
        draw.rounded_rectangle((x0, y, x0 + width, y + 65), radius=12, fill=color)
        draw.text((x0 + width + 22, y + 15), f"{value:.1f}s",
                  font=_font(27, True), fill="#0B2545")
    draw.text((1110, 485), "85.5% reduction | 6.9x faster",
              font=_font(25, True), fill="#15803D")
    path1 = WORK / "resolver_latency.png"
    image.save(path1, quality=95)

    # Chart 2: baseline measured stage breakdown vs projected optimized run.
    image, draw = _chart_canvas(1600, 700, "Where the end-to-end time went")
    stages = ["Schedule + extraction", "Unclaimed agent wait",
              "Index + cold setup", "Evidence resolution", "Finalize"]
    before = [34.3, 30.1, 31.4, 157.5, 1.2]
    optimized = [34.3, 8.0, 17.0, 22.8, 1.2]
    palette = ["#64748B", "#B98900", "#2A9D8F", "#2E74B5", "#94A3B8"]
    x0, x1, max_v = 330, 1480, 280
    for tick in (0, 50, 100, 150, 200, 250):
        x = x0 + int((x1 - x0) * tick / max_v)
        draw.line((x, 135, x, 455), fill="#E2E8F0", width=2)
        draw.text((x - 13, 470), str(tick), font=_font(20), fill="#64748B")
    for y, label, values in ((190, "Measured baseline", before),
                             (345, "Optimized projection", optimized)):
        draw.text((55, y + 17), label, font=_font(24, True), fill="#334155")
        left = x0
        for value, color in zip(values, palette):
            width = max(4, int((x1 - x0) * value / max_v))
            draw.rectangle((left, y, left + width, y + 72), fill=color)
            if width > 68:
                txt = f"{value:.1f}"
                box = draw.textbbox((0, 0), txt, font=_font(21, True))
                draw.text((left + (width - (box[2] - box[0]))/2, y + 22), txt,
                          font=_font(21, True), fill="white")
            left += width
        draw.text((left + 18, y + 18), f"{sum(values):.1f}s",
                  font=_font(25, True), fill="#0B2545")
    draw.text((650, 505), "End-to-end seconds", font=_font(21), fill="#64748B")
    legend_x, legend_y = 55, 570
    for idx, (stage, color) in enumerate(zip(stages, palette)):
        col = idx % 3
        row = idx // 3
        x = legend_x + col * 505
        y = legend_y + row * 45
        draw.rounded_rectangle((x, y, x + 25, y + 25), radius=4, fill=color)
        draw.text((x + 37, y - 1), stage, font=_font(20), fill="#475569")
    path2 = WORK / "end_to_end_stages.png"
    image.save(path2, quality=95)

    # Chart 3: frozen accuracy metrics.
    image, draw = _chart_canvas(1600, 760, "Frozen holdout accuracy: baseline vs optimized")
    labels = ["Top-1", "Recall@3", "Recall@5", "Recall@10", "MRR"]
    baseline = [85.8824, 94.1176, 96.4706, 97.0588, 90.1816]
    after = [85.8824, 94.1176, 96.4706, 97.0588, 90.1599]
    xmin, xmax, x0, x1 = 84.0, 99.5, 350, 1450
    for tick in (84, 87, 90, 93, 96, 99):
        x = x0 + int((x1 - x0) * (tick - xmin) / (xmax - xmin))
        draw.line((x, 125, x, 625), fill="#E2E8F0", width=2)
        draw.text((x - 21, 650), f"{tick}%", font=_font(20), fill="#64748B")
    for idx, (label, base, current) in enumerate(zip(labels, baseline, after)):
        y = 165 + idx * 98
        draw.text((80, y - 14), label, font=_font(25, True), fill="#334155")
        xb = x0 + int((x1 - x0) * (base - xmin) / (xmax - xmin))
        xc = x0 + int((x1 - x0) * (current - xmin) / (xmax - xmin))
        draw.line((min(xb, xc), y, max(xb, xc), y), fill="#94A3B8", width=5)
        draw.ellipse((xb - 12, y - 12, xb + 12, y + 12), fill="white",
                     outline="#64748B", width=4)
        draw.ellipse((xc - 8, y - 8, xc + 8, y + 8), fill="#2E74B5")
        draw.text((max(xb, xc) + 22, y - 15), f"{current:.2f}%",
                  font=_font(22, True), fill="#0B2545")
    draw.ellipse((1040, 690, 1064, 714), fill="white", outline="#64748B", width=4)
    draw.text((1078, 686), "Frozen baseline", font=_font(20), fill="#475569")
    draw.ellipse((1280, 694, 1296, 710), fill="#2E74B5")
    draw.text((1310, 686), "Optimized", font=_font(20), fill="#475569")
    path3 = WORK / "accuracy_holdout.png"
    image.save(path3, quality=95)
    return path1, path2, path3


def build():
    WORK.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    chart1, chart2, chart3 = create_charts()

    doc = Document()
    doc.core_properties.title = "VEDA Latency Optimization Report"
    doc.core_properties.subject = "Measured performance improvement and accuracy assurance"
    doc.core_properties.author = "Shivansh Mukhia"
    doc.core_properties.keywords = "VEDA, latency, MetaRank, Rescheduler, performance, accuracy"
    doc.core_properties.comments = "Prepared from measured VEDA runtime and frozen holdout results."
    doc.settings.odd_and_even_pages_header_footer = False
    set_document_styles(doc)
    for section in doc.sections:
        configure_section(section)
        add_running_furniture(section)

    # Page 1 - editorial cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("PERFORMANCE ENGINEERING REPORT"), size=10.5,
            bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("VEDA Latency Optimization"), size=30,
            bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Measured speedup, design rationale and accuracy assurance"),
            size=14, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(26)
    set_paragraph_border_bottom(rule, color=BLUE, size=14, space=4)

    metrics = doc.add_table(rows=2, cols=3)
    set_table_geometry(metrics, [3120, 3120, 3120], indent=90)
    set_repeat_table_header(metrics.rows[0])
    set_table_borders(metrics, color="DCE5EE", size=5)
    metric_values = [("6.9x", "Faster resolver"),
                     ("85.5%", "Latency reduction"),
                     ("0.00 pp", "Top-1 accuracy change")]
    for idx, (value, label) in enumerate(metric_values):
        set_cell_shading(metrics.cell(0, idx), PALE_BLUE if idx < 2 else PALE_GREEN)
        set_cell_shading(metrics.cell(1, idx), WHITE)
        pv = metrics.cell(0, idx).paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pv.paragraph_format.space_after = Pt(0)
        set_run(pv.add_run(value), size=21, bold=True,
                color=BLUE if idx < 2 else GREEN)
        pl = metrics.cell(1, idx).paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_after = Pt(0)
        set_run(pl.add_run(label), size=9.2, bold=True, color=NAVY)
        for cell in (metrics.cell(0, idx), metrics.cell(1, idx)):
            set_cell_margins(cell, top=140, bottom=140, start=90, end=90)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Prepared by Shivansh Mukhia"), size=12,
            bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("28 August 2026 | VEDA project performance review"),
            size=9.5, color=MUTED)

    doc.add_page_break()

    # Page 2 - executive results.
    add_heading(doc, "1. Executive result", 1)
    add_callout(
        doc,
        "CONCLUSION",
        "The optimized resolver processed the same 165 evidence records in 22.8 seconds, down from approximately 157.5 seconds. The frozen 170-case accuracy benchmark retained exactly the same Top-1 and recall scores.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_text(
        doc,
        "The original 254.5-second end-to-end run was not uniformly slow. Most of the avoidable delay came from repeated candidate reasoning, cold machine-learning imports and an unclaimed Antigravity inbox window. The improvement therefore targeted those measured bottlenecks instead of weakening the resolver.",
    )
    add_figure(
        doc, chart1,
        "Figure 1. Resolver latency on the real 165-record project.",
        "Horizontal bars compare 157.5 seconds before optimization with 22.8 seconds after optimization, an 85.5 percent reduction.",
    )
    add_heading(doc, "Measured comparison", 2)
    add_table(
        doc,
        ["Metric", "Before", "After", "Change"],
        [
            ["Resolver time (165 records)", "157.5 s", "22.8 s", "-85.5%"],
            ["Resolver throughput", "1.05 records/s", "7.24 records/s", "6.9x"],
            ["Per-record resolver time", "0.955 s", "0.138 s", "-85.5%"],
            ["Embedding selector cold start", "14.7 s", "0.26 s", "-98.2%"],
            ["Unclaimed Antigravity wait", "30 s", "8 s", "-73.3%"],
            ["Full analysis", "254.5 s measured", "60-90 s expected", "~67% midpoint"],
        ],
        [4200, 1750, 1750, 1660],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_text(
        doc,
        "Measurement note: the 22.8-second resolver result is directly measured. The optimized full-analysis range is an engineering estimate because Horizun startup and claimed Antigravity inference time vary by machine and run.",
        size=9.2, color=MUTED, after=0,
    )

    doc.add_page_break()

    # Page 3 - root cause and approach.
    add_heading(doc, "2. Root cause and optimization approach", 1)
    add_text(
        doc,
        "Durable job timestamps separated the run into schedule ingestion, provider handoff, index preparation, evidence resolution and final validation. This showed that UI rendering was not the main latency source: repeated backend work was.",
    )
    add_figure(
        doc, chart2,
        "Figure 2. Measured baseline stage breakdown and projected optimized first run.",
        "Stacked horizontal bars show the baseline total of 254.5 seconds and an optimized projection of 83.3 seconds across schedule extraction, provider wait, index setup, evidence resolution and finalization.",
    )
    add_heading(doc, "Why these changes were selected", 2)
    add_table(
        doc,
        ["Bottleneck", "Optimization", "Why this approach"],
        [
            ["Cold hardware detection", "Lightweight GPU probe", "Avoid importing PyTorch only to choose the dependency-free CPU backend."],
            ["Cold MetaRank load", "Background model warmup", "Moves XGBoost and frozen-model loading outside the first evidence request."],
            ["Repeated feature work", "Cache invariant schedule and evidence context", "Computes snapshot, corroboration and logic mode once per row instead of once per candidate."],
            ["Large object copies", "Shallow ranking envelopes", "Copies only mutable scores and features; immutable activity documents are shared safely."],
            ["Rescheduler hot loop", "Precompute dates and schedule state", "Removes tens of thousands of repeated parses and status checks without changing beam depth or equations."],
            ["Idle provider bridge", "Eight-second claim guard", "Releases the worker quickly when no Antigravity agent is polling; a claimed job still gets the normal reasoning window."],
        ],
        [1900, 2600, 4860],
        font_size=8.7,
    )
    add_callout(
        doc,
        "DESIGN PRINCIPLE",
        "Make the same decision with less repeated work. No expert was removed, no candidate list was shortened, and no validation or schedule-write gate was bypassed.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    doc.add_page_break()

    # Page 4 - accuracy and limits.
    add_heading(doc, "3. Accuracy impact", 1)
    add_callout(
        doc,
        "ANSWER",
        "No material accuracy loss was observed. Top-1 and Recall@3/5/10 were identical across all 170 frozen holdout cases.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_figure(
        doc, chart3,
        "Figure 3. Frozen holdout ranking quality before and after optimization.",
        "Dot plot compares baseline and optimized Top-1, Recall at 3, Recall at 5, Recall at 10 and mean reciprocal rank across 170 labeled cases.",
        width=5.65,
    )
    add_table(
        doc,
        ["Accuracy metric", "Frozen baseline", "Optimized", "Delta"],
        [
            ["Top-1 accuracy", "85.88%", "85.88%", "0.00 pp"],
            ["Recall@3", "94.12%", "94.12%", "0.00 pp"],
            ["Recall@5", "96.47%", "96.47%", "0.00 pp"],
            ["Recall@10", "97.06%", "97.06%", "0.00 pp"],
            ["Mean Reciprocal Rank", "90.1816%", "90.1599%", "-0.0218 pp"],
            ["Expert failures", "Not applicable", "0 of 170", "None"],
        ],
        [3300, 2020, 2020, 2020],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.8,
        cell_tb=55,
    )
    add_text(
        doc,
        "MRR shifted by just -0.0218 percentage points, affecting only ordering below the primary decision boundary; Top-1 and recall were unchanged.",
        size=9.2,
    )
    add_heading(doc, "Why accuracy was preserved", 2)
    add_text(
        doc,
        "The production path still creates Semantic, Engineering, Tree and Rescheduler-v2 lists, fuses them with the same frozen LambdaMART models, then applies calibration, deterministic validators and risk policy. Only repeated input preparation changed.",
        size=10.2,
    )
    add_callout(
        doc,
        "NEXT BOTTLENECK",
        "Horizun/Microsoft Project startup remains a roughly 28-second fixed cost. Keep the schedule service warm and stream linked evidence in small batches. Evidence: durable job timestamps, the real 165-record benchmark, a 170-case frozen holdout with 13,436 activities, and six regression suites.",
        fill=LIGHT,
        accent=DARK_BLUE,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
